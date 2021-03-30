# -*- coding: utf-8 -*-
"""
Created on Wed Mar 17 11:23:32 2021

@author: raha_
"""
#text canonization
from docx import Document
import re
import spacy
from nltk import sent_tokenize
from docx.enum.style import WD_STYLE_TYPE
from num2words import num2words
import json
from collections import Counter


#-----------------------------------------------------------------
#find acronyoms and thier expansion and store them in a dictionary
def expand_abv(Paras):
    #dictionary for those acrs which has an extra word    
    arts = ['a', 'an', 'and', 'the', 'of', 'in', 'with', 'by', 'from']
   #save acrs
    acr={}
    for p in Paras:
        for run in p.runs:
            t=run.text 
            #removing hyphens
            puncts='_-—'
            for sym in puncts:
                t= t.replace(sym,' ')  
                #find acrs (ABCD)
            mo=re.findall(r'\([A-Z]*\)', t) 
            #print(mo)
             
            words = t.split(' ') 
            #Expand the acr, find Ngrams prior to (ABC)
            if mo:                                 
                for item in mo:                    
                    if item.isupper():                    
                        l=len(item)-1 
                        #print(l)
                        ng=''
                        for i, j in enumerate(words):                        
                            if j == item:
                                artno=0
                                for k in range(1,l):
                                    w=words[i-k]
                                    if w in arts:
                                        artno+=1
                                    ng+=w
                                    ng+=' '
                                #of the ,... management
                                if artno>0:
                                    for c in range(artno):                                    
                                        w=words[i-(k+c+1)]
                                        ng+=w
                                        ng+=' '
                                #reverse the phrase
                                sng=ng.split()        
                                l1= len(sng) 
                                rng=''
                                for i in reversed(range(l1)):
                                    rng+=sng[i]
                                    rng+=' '                               
                                    #add to dictionary
                                
                                acr[item]=rng
    
    facr={}
    #print(json.dumps(acr, indent = 4))
    for item, rng in acr.items():
        i=item.replace('(','')
        i=i.replace(')','')
        facr[i]=rng
        
    #print(json.dumps(facr, indent = 4))
    return facr

    
               
#-----------------------------------------------------------------
#convert numbers and dates to strings using regex
def convert_num(text):

    timer=re.findall(r'\d+:\d+[ am| AM| pm| PM]', text)   
    for t in timer:
        t1=t
        t=t.replace(':',' ')
        text= text.replace(t1,t)
    timer=re.findall(r'\d+:\d+[am|AM|pm|PM]', text)   
    for t in timer:
        text=t.replace(':',' ')   
        text= text.replace(time,time1)   
    mo=re.findall(r'\d+', text)    
    if mo:        
        for item in mo:
            sa=num2words(item)  
            text = re.sub(r'\d+',sa,text)          
    return text
        

 
#-----------------------------------------------------------------
#removing symbolls and runs
def clean_text(text):
    
    puncts='_-()"”“©—*:\''
    for sym in puncts:
        text= text.replace(sym,' ')        
    reps='%&/#+'
    for sym in reps:
        if sym=='%':
            text= text.replace(sym,' percent ')
        if sym=='&':
            text= text.replace(sym,' and ')
        if sym=='/':
            text= text.replace(sym,' or ')
        if sym=='#':
            text= text.replace(sym,' number ')
        if sym=='+':
            text= text.replace(sym,' plus ')
        if sym=='=':
            text= text.replace(sym,' equals ')
    return text

#-----------------------------------------------------------------
#Remove country abriviations
def Country_abv(text):    
    acronyms = [
        ('U.S.', 'United States'), 
        ('USA', 'United States'),
        ('US', 'United States'),
        ('U.S.', 'United States'),
        ('UK', 'United Kingdom'),
        ('U.K.', 'United Kingdom'),
        ('Great Britain', 'United Kingdom'),
        ('Britain', 'United Kingdom'),        
    ]
    pattern = '|'.join('(%s)' % re.escape(match) for match, replacement in acronyms)
    substitutions = [match for replacement, match in acronyms]
    replace = lambda m: substitutions[m.lastindex - 1]
    return re.sub(pattern, replace, text)

#-----------------------------------------------------------------
def lang_abv(text):    
    acronyms = [
        ('i.e.', 'for example'),
        ('e.g.', 'for instance') ,      
        ('’ve',' have' ),
        ('’re', ' are'),        
        ('I’m',' I am' ),
        ('n’t', ' not'),
        ('’s', ' is')
        
        
    ]
    
    pattern = '|'.join('(%s)' % re.escape(match) for match, replacement in acronyms)
    substitutions = [match for replacement, match in acronyms]
    replace = lambda m: substitutions[m.lastindex - 1]
    
    return re.sub(pattern, replace, text)

#-----------------------------------------------------------------
#expand the acrs
def acrs(text,da): 
    if da:
        #convert dict of acrss to list
        acronyms=[]
        temp=[]
        for key, value in da.items():
            temp = [key,value]
            acronyms.append(temp) 
         
        
        pattern = '|'.join('(%s)' % re.escape(match) for match, replacement in acronyms) 
        substitutions = [match for replacement, match in acronyms]
        replace = lambda m: substitutions[m.lastindex - 1]
        ft=re.sub(pattern, replace, text)
    
        return ft,acronyms
    else: return text, []

#-----------------------------------------------------------------
#convert money symbols
def get_text(doc):   
    fullText = []
    for para in doc:
        fullText.append(para.text)
    return '\n'.join(fullText)  


#-----------------------------------------------------------------
#replace currency values such as $ with their respective text
def cur_remove(text):
    t=''
    sp=text.split()
    puncts='€$,.;'
    for w in sp: 
        #print(w)
        if w.startswith('$'):
           # print(w)
            w+=' Dollar '            
            for p in puncts:
                if p=='$':
                    w=w.replace(p, "")
                    #print(w)                  
            
        
        if w.startswith('€'):
            w+=' Euro '
            for p in puncts:
                if p=='€':
                    w=w.replace(p, "")                          
         
        t=t +' ' +w        
    return t

  
#-----------------------------------------------------------------
#concatenate lists and bullet points and removing empty line
def file_creation(doc):    
    Paras = normDoc.paragraphs   
    l=len(Paras) 
    fnormDoc= Document()
    p1=normDoc.paragraphs[0] 
    #print(l)
    st=str(p1.text)
       #print(st)
    fnormDoc.add_paragraph(st)
    #connct succesive lines
    i=0
    while i<l-1:
       i=i+1    
       #print(i)
       p1=normDoc.paragraphs[i] 
       l2=len(p1.text)       
       if l2<50:           
           for j in  range(i+1,l):
               p=normDoc.paragraphs[j] 
               l1=len(p.text) 
               #jump after the list items
               if l1>=50: 
                   i=j
                   p1.add_run(p.text)                   
                   break          
                   
               else:
                   p1.add_run(p.text) 
                   p1.add_run(',')
                   i=j
                   
       else:
           p1=normDoc.paragraphs[i]
       #print(p1.text)
       st=str(p1.text)
       #print(st)
       fnormDoc.add_paragraph(st)
       fnormDoc.save("11.docx")
       
#-----------------------------------------------------------------
#flatten the lists
def list_norm(pars):   
    noList= Document()

    for p in Paras:
        for run in p.runs:
            x=re.sub(r'^\w[.)]\w*', '', run.text) 
            noList.add_paragraph(x)
    
    return noList

#-----------------------------------------------------------------
#removing all the styles
def remove_styles(doc):
        
    st=doc.styles
    for s in st:    
        strs=s.name        
        st[strs].delete
    return doc

#-----------------------------------------------------------------
#remove the (ACR) to prevent duplicate words
def remove_first(Text):
    
    mo=re.findall(r'\([A-Z]*\)', Text) 
        #print(mo)
    if mo:
        for item in mo: 
            Text=Text.replace(item,'')
    
    return(Text)
                   
    
           
        
    

#main code

doc = Document('bench1_un.docx')
normDoc= Document()
Paras = doc.paragraphs
#create a list of acronyms
DictAcr=expand_abv(Paras)
#removing lists
docwl = list_norm(Paras)
#deleting styles
doc =remove_styles(docwl)

Paras=doc.paragraphs
for p in Paras:    
    for run in p.runs:        
        T=run.text 
        #remove ACR instance from text eg genetic alg (ga)> genetic alg
        T=remove_first(T)
        #resolve and expand all the Acronyms
        [T,L]=acrs(T,DictAcr)   
        #replace currency abv
        CurPNCC=cur_remove(T)
        #convert dates and nums to text
        Cur=convert_num(CurPNCC) 
        #resolve counrtry abvs        
        CurPNC=Country_abv(Cur) 
        #resolve other abvs
        CurL=lang_abv(CurPNC)            
        CurPN=clean_text(CurL)                         
        normDoc.add_paragraph(CurPN)
        #normDoc=re.sub(r'\s*\n\s*', ' ', normDoc)
     
file_creation(normDoc)

       


    
    
       
    
        
        
       
  
        
        
        


